"""
file_engine/word_writer.py
===========================
สร้างไฟล์ Word (.docx) รวมภาพ screenshot ทั้งหมด — เลียนแบบโครงสร้างของไฟล์
ตัวอย่างจริงของผู้ใช้ (Maize_C425201_R450006SandLo 231068.docx, 1,204 ภาพ)
ที่ยืนยันด้วยการแกะ document.xml แล้วว่าเรียบง่ายมาก:

    ต่อ 1 วันปลูก:  บรรทัดวันที่ "1/4/1981" (d/m/yyyy ไม่มีเลขศูนย์นำ)
                    ตามด้วยรูป 2 ใบ: ตาราง schedule แล้วก็กราฟ
    เรียงตามปฏิทิน ปีเก่า -> ใหม่, ในปีเรียงตามวันที่ปลูก — ไม่มีหัวข้อ/สารบัญอื่นเลย

    ขนาดหน้า: A4 แนวตั้ง, รูปกว้าง ~7.27 นิ้ว (เต็มความกว้างหน้าแบบขอบแคบ)

เหมือนหลักการเดียวกับ excel_writer: อ่านจากไฟล์ภาพที่มีอยู่จริงในโฟลเดอร์
screenshots ณ ตอนนั้น (สร้างโดย capture_screenshots ชื่อ {ปี}_{MMDD}_schedule.png
/ _graph.png) — เรียกซ้ำได้ทุกเมื่อ เขียนทับไฟล์ Word เดิมทั้งไฟล์ ไม่ผูกกับ
state การรันของเฟส 1 เลย
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger("word_writer")

# ชื่อไฟล์ภาพที่ capture_screenshots สร้าง: "{ปี}_{MMDD}_schedule.png"
SCREENSHOT_RE = re.compile(r"^(?P<year>\d{4})_(?P<mmdd>\d{4})_schedule\.png$")

IMAGE_WIDTH_INCHES = 7.27  # วัดจากไฟล์ตัวอย่างจริง (wp:extent 6645910 EMU)


def build_screenshot_doc(screenshot_dir: Path, output_docx: Path) -> int:
    """สร้าง .docx จากภาพทั้งหมดใน screenshot_dir คืนจำนวนวันปลูกที่ใส่ลงเอกสาร
    (import python-docx ในฟังก์ชันเพื่อไม่ให้กระทบ startup ถ้าไลบรารีมีปัญหา)"""
    from docx import Document
    from docx.shared import Inches, Mm

    screenshot_dir = Path(screenshot_dir)
    if not screenshot_dir.is_dir():
        raise FileNotFoundError(f"ไม่พบโฟลเดอร์ภาพ: {screenshot_dir}")

    # จับคู่ (ปี, วันปลูก) -> (ภาพตาราง, ภาพกราฟ) เรียงตามปฏิทิน
    entries: list[tuple[int, str, Path, Path]] = []
    for schedule_path in sorted(screenshot_dir.glob("*_schedule.png")):
        m = SCREENSHOT_RE.match(schedule_path.name)
        if not m:
            logger.warning("ข้ามไฟล์ชื่อไม่ตรงรูปแบบที่คาด: %s", schedule_path.name)
            continue
        graph_path = schedule_path.with_name(schedule_path.name.replace("_schedule", "_graph"))
        entries.append((int(m.group("year")), m.group("mmdd"), schedule_path, graph_path))
    entries.sort(key=lambda e: (e[0], e[1]))

    if not entries:
        raise FileNotFoundError(f"ไม่พบภาพ screenshot เลยในโฟลเดอร์: {screenshot_dir}")

    doc = Document()
    # A4 แนวตั้ง ขอบแคบ ให้รูปกว้าง 7.27 นิ้วพอดีแบบไฟล์ตัวอย่าง
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = section.right_margin = Mm(12.7)
    section.top_margin = section.bottom_margin = Mm(12.7)

    count = 0
    for year, mmdd, schedule_path, graph_path in entries:
        month, day = int(mmdd[:2]), int(mmdd[2:])
        # รูปแบบวันที่ตามไฟล์ตัวอย่างเป๊ะ: "1/4/1981" (d/m/yyyy ไม่มีเลขศูนย์นำ)
        doc.add_paragraph(f"{day}/{month}/{year}")
        doc.add_picture(str(schedule_path), width=Inches(IMAGE_WIDTH_INCHES))
        if graph_path.exists():
            doc.add_picture(str(graph_path), width=Inches(IMAGE_WIDTH_INCHES))
        else:
            logger.warning("ไม่พบภาพกราฟคู่ของ %s (%s)", schedule_path.name, graph_path.name)
        count += 1

    output_docx = Path(output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return count
